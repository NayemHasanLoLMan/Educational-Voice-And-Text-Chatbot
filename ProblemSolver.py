# import openai
# import pytesseract
# from pdfminer.high_level import extract_text as extract_pdf_text
# from docx import Document
# from PIL import Image
# import os
# from dotenv import load_dotenv

# load_dotenv()

# def __init__ (content, api_key: str = None, language: str = "oromo"):
# # Load environment variables
#     if api_key:
#         openai.api_key = api_key
#     else:
#         openai.api_key = os.getenv("OPENAI_API_KEY")


#     #new section for language
#     language = language.lower()
#     supported_languages = {"english", "amharic", "oromo"}
#     if language not in supported_languages:
#         language = "english"

#     language = language



#     if language == "oromo":

#         language_instruction = "Please ONLY use **Oromo** throughtout the entire process of creating the solution and for the solution"

#     if language =="amharic":

#         language_instruction="Please ONLY use **Amharic** throughtout the entire process of creating the solution and for the solution"

#     else:

#         language_instruction = "Please ONLY use **English** throughtout the entire process of creating the solution and for the solution"

#     prompt = f"""
#         You are an expert AI tutor and analyst. Analyze the provided text and follow any instructions or context provided in the user input to deliver a comprehensive response. Your goal is to understand the content, identify key elements, and provide solutions, explanations, or insights as directed.

#         ### Content to Analyze:
#         {content}

#         ### User Input:
#         {user_input if user_input else 'No user input provided. Analyze the content, identify any questions, problems, or key topics, and provide comprehensive solutions or explanations.'}

#         ### Analysis Guidelines:
#         1. **Content Understanding**:
#         - Summarize the main topics, themes, or purpose of the document.
#         - Identify the type of document (e.g., educational, technical, narrative, question-based, etc.).
#         - Highlight any questions, problems, exercises, or key concepts present.

#         2. **Response Requirements**:
#         - If the user input contains instructions, follow them explicitly (e.g., solve problems, explain concepts, analyze issues, etc.).
#         - If the user input includes additional questions or content, incorporate them into the analysis.
#         - For questions or problems (from the document or user input):
#             - Extract and list all questions, exercises, or problems.
#             - Provide step-by-step solutions or explanations.
#             - Include relevant formulas, theories, or concepts with clear explanations.
#             - Use examples or analogies relevant to the context (e.g., Ethiopian culture for educational content).
#             - Highlight key points and common pitfalls.
#         - For general analysis (if no questions or specific instructions):
#             - Discuss key insights, implications, or applications of the content.
#             - Suggest improvements, alternative approaches, or further exploration if relevant.
#         - If user input specifies a method or focus (e.g., 'solve using a specific formula' or 'focus on cultural relevance'), prioritize that approach.

#         3. **Content-Specific Handling**:
#         - **Mathematical Problems**: Show complete calculations, explain each step, and include relevant formulas.
#         - **Theoretical Questions**: Provide detailed explanations with examples and context.
#         - **Multiple Choice Questions**: Identify the correct answer and explain why other options are incorrect.
#         - **Open-Ended Questions**: Explore multiple perspectives, provide balanced arguments, and support with examples.
#         - **Code-Related Content**: Include code snippets, explain functionality, and highlight best practices.
#         - **Non-Question Content**: Analyze the document's purpose, structure, and key points; provide insights or recommendations.

#         4. **Formatting**:
#         - Use clear headers (e.g., 'Document Summary', 'Identified Questions', 'Solutions', 'Key Insights').
#         - Use bullet points, numbered steps, or tables for clarity.
#         - Ensure readability for a broad audience, including young learners if the context suggests it (e.g., Ethiopian students in grades 1-12).
#         - Include culturally relevant examples or context where applicable.

#         ### Deliverables:
#         - A comprehensive response addressing the user input and content analysis.
#         - Clear, concise, and well-structured explanations or solutions.
#         - Practical examples or illustrations to enhance understanding.
#         - Suggestions for further learning or application if relevant.

#         Format your response in Markdown with clear headers, bullet points, and numbered steps for better readability.
#         """
        
        
#     system_prompt = language_instruction + "\n\n" + prompt

# #############################################

# def extract_text_from_file(file_path):
#     """Extracts text from PDF, DOCX, or image files."""
#     file_extension = file_path.split('.')[-1].lower()

#     if file_extension == "pdf":
#         return extract_pdf_text(file_path)

#     elif file_extension == "docx":
#         doc = Document(file_path)
#         return "\n".join([para.text for para in doc.paragraphs])

#     elif file_extension in ["png", "jpg", "jpeg"]:
#         image = Image.open(file_path)
#         return pytesseract.image_to_string(image)

#     else:
#         raise ValueError("Unsupported file format. Please upload a PDF, DOCX, or image.")

# def generate_text(system_prompt):
#     """Generates text using OpenAI's GPT model based on the provided system_prompt."""
#     response = openai.ChatCompletion.create(
#         model="gpt-4-turbo",
#         messages=[{"role": "user", "content": system_prompt}],
#         max_tokens=3000
#     )
#     return response.choices[0].message.content.strip()

# def solve_questions(system_prompt, user_input=None):
#     """Analyzes content and provides comprehensive solutions or insights based on user input."""

    
#     return generate_text(system_prompt)

# def create_ai_step_by_step_solution(file_path=None, user_input=None):
#     """Processes the uploaded file and/or user input, extracts text, and generates a comprehensive solution or analysis."""
#     try:
#         combined_content = ""

#         # Extract text from file if provided
#         if file_path:
#             if not os.path.exists(file_path):
#                 print(f"Error: File '{file_path}' not found.")
#                 return
#             print(f"Extracting text from '{file_path}'...\n")
#             file_content = extract_text_from_file(file_path)
#             if file_content.strip():
#                 combined_content += f"--- Content from File ---\n{file_content}\n"
#             else:
#                 print("No text detected in the file. Checking user input...")

#         # Append user input if provided
#         if user_input:
#             if user_input.strip():
#                 combined_content += f"--- User Input ---\n{user_input}\n"
#             else:
#                 print("No valid user input provided.")

#         # Check if there's any content to process
#         if not combined_content.strip():
#             print("No valid content detected from file or user input. Please provide a file or input.")
#             return

#         print("\n--- Extracted Content Preview ---\n")
#         print(combined_content[:1000] + "...")  # Show first 1000 characters for preview

#         print("\n--- AI Comprehensive Analysis and Solutions ---\n")
#         solutions = solve_questions(combined_content, user_input)
        
#         print(solutions)
        
#     except Exception as e:
#         print(f"An error occurred: {str(e)}")

# Example usage
# if __name__ == "__main__":
#     file_path = "C:\\Users\\hasan\\Downloads\\12th_grade_math_problems.pdf"  # Replace with actual file path
#     user_input = """
#     analyze the document and solve the 12 grade problems in the pdf step by step and give me the solution in a well structured format.
#     """
#     create_ai_step_by_step_solution(file_path=file_path, user_input=user_input)


import openai
import pytesseract
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document
from PIL import Image
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class DocumentAnalyzer:
    def __init__(self, api_key: str = None, language: str = None):
        """
        Initialize the DocumentAnalyzer with API key and language settings.
        
        Args:
            api_key (str, optional): OpenAI API key. If None, loads from environment.
            language (str): Response language ('english', 'amharic', 'oromo')
        """
        # Set up OpenAI API key
        if api_key:
            openai.api_key = api_key
        else:
            api_key_from_env = os.getenv("OPENAI_API_KEY")
            if not api_key_from_env:
                raise ValueError("OpenAI API key not found. Please provide it or set OPENAI_API_KEY environment variable.")
            openai.api_key = api_key_from_env
        
        # Set language
        self.language = language.lower()
        supported_languages = {"english", "amharic", "oromo"}
        if self.language not in supported_languages:
            print(f"Warning: Language '{language}' not supported. Defaulting to 'english'.")
            self.language = "english"
    
    def extract_text_from_file(self, file_path):
        """
        Extracts text from PDF, DOCX, or image files.
        
        Args:
            file_path (str): Path to the file to extract text from
            
        Returns:
            str: Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File '{file_path}' not found.")
        
        file_extension = file_path.split('.')[-1].lower()
        
        try:
            if file_extension == "pdf":
                return extract_pdf_text(file_path)
            
            elif file_extension == "docx":
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            
            elif file_extension in ["png", "jpg", "jpeg", "bmp", "tiff"]:
                image = Image.open(file_path)
                return pytesseract.image_to_string(image)
            
            else:
                raise ValueError(f"Unsupported file format: {file_extension}. Please upload a PDF, DOCX, or image file.")
                
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")
    
    def _create_analysis_prompt(self, content, user_input=None):
        """
        Create the analysis prompt with content and user input.
        
        Args:
            content (str): Extracted content to analyze
            user_input (str, optional): Additional user instructions
            
        Returns:
            str: Complete system prompt for AI analysis
        """
        # Language instruction - only for response format
        if self.language == "oromo":
            language_instruction = "Please provide your entire response in Oromo language only."
        elif self.language == "amharic":
            language_instruction = "Please provide your entire response in Amharic language only."
        else:
            language_instruction = "Please provide your entire response in English language only."
        
        prompt = f"""
        {language_instruction}

        You are an expert AI tutor and analyst. Analyze the provided text and follow any instructions or context provided in the user input to deliver a comprehensive response.

        ### Content to Analyze:
        {content}

        ### User Input:
        {user_input if user_input else 'No user input provided. Analyze the content, identify any questions, problems, or key topics, and provide comprehensive solutions or explanations.'}

        ### Analysis Guidelines:
        1. **Content Understanding**:
        - Summarize the main topics, themes, or purpose of the document.
        - Identify the type of document (e.g., educational, technical, narrative, question-based, etc.).
        - Highlight any questions, problems, exercises, or key concepts present.

        2. **Response Requirements**:
        - If the user input contains instructions, follow them explicitly.
        - For questions or problems: Extract and list all questions, provide step-by-step solutions.
        - Include relevant formulas, theories, or concepts with clear explanations.
        - For general analysis: Discuss key insights, implications, or applications.

        3. **Content-Specific Handling**:
        - **Mathematical Problems**: Show complete calculations and explain each step.
        - **Theoretical Questions**: Provide detailed explanations with examples.
        - **Multiple Choice Questions**: Identify correct answers and explain why others are incorrect.
        - **Open-Ended Questions**: Explore multiple perspectives with balanced arguments.

        4. **Formatting**:
        - Use clear headers and bullet points for clarity.
        - Ensure readability for a broad audience.
        - Format your response with clear structure and numbered steps.

        Provide a comprehensive response addressing the user input and content analysis with clear, well-structured explanations or solutions.
        """
        
        return prompt
    
    def generate_analysis(self, system_prompt):
        """
        Generates analysis using OpenAI's GPT model.
        
        Args:
            system_prompt (str): The complete prompt for analysis
            
        Returns:
            str: Generated analysis from the AI model
        """
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": system_prompt}],
                max_tokens=3000,
                temperature=0.7
            )
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            raise Exception(f"Error generating analysis: {str(e)}")
    
    def analyze_content(self, content, user_input=None):
        """
        Analyzes content and provides solutions in the selected language.
        
        Args:
            content (str): Content to analyze
            user_input (str, optional): Additional user instructions
            
        Returns:
            str: AI-generated analysis in selected language
        """
        system_prompt = self._create_analysis_prompt(content, user_input)
        return self.generate_analysis(system_prompt)
    
    def process_document(self, file_path=None, user_input=None):
        """
        Processes document and generates analysis in selected language.
        
        Args:
            file_path (str, optional): Path to the document file
            user_input (str, optional): Additional user instructions
            
        Returns:
            str: Complete analysis in selected language
        """
        try:
            combined_content = ""
            
            # Extract text from file if provided
            if file_path:
                print(f"Extracting text from '{file_path}'...")
                file_content = self.extract_text_from_file(file_path)
                if file_content.strip():
                    combined_content += f"--- Content from File ---\n{file_content}\n"
                    print("✓ Text extracted successfully")
                else:
                    print("⚠ No text detected in file")
            
            # Add user input if provided
            if user_input and user_input.strip():
                combined_content += f"--- User Input ---\n{user_input}\n"
                print("✓ User input added")
            
            # Check if there's content to process
            if not combined_content.strip():
                raise ValueError("No content provided. Please provide a file or user input.")
            
            print(f"\n--- Generating Analysis in {self.language.title()} ---")
            analysis = self.analyze_content(combined_content, user_input)
            
            return analysis
            
        except Exception as e:
            print(f"❌ Error: {str(e)}")
            return None

def create_solution(file_path=None, user_input=None, language: str = None):
    """
    Simple function to analyze document and get response in selected language.
    
    Args:
        file_path (str, optional): Path to document
        user_input (str, optional): User instructions
        language (str): Response language ('english', 'amharic', 'oromo')
    """
    try:
        analyzer = DocumentAnalyzer(language=language)
        result = analyzer.process_document(file_path=file_path, user_input=user_input)
        
        if result:
            print("\n" + "="*50)
            print(f"ANALYSIS RESULT ({language.upper()})")
            print("="*50)
            print(result)
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")

# Example usage
if __name__ == "__main__":
    # Example with English response
    create_solution(
        file_path="C:\\Users\\hasan\\Downloads\\12th_grade_math_problems.pdf",
        user_input="Solve the math problems step by step",
        language="english"
    )
    
    # Example with Amharic response
    # create_solution(
    #     file_path="C:\\Users\\hasan\\Downloads\\12th_grade_math_problems.pdf",
    #     user_input="Solve the math problems step by step",
    #     language="amharic"
    # )
    
    #Example with Oromo response
    # create_solution(
    #     file_path="C:\\Users\\hasan\\Downloads\\12th_grade_math_problems.pdf",
    #     user_input="Solve the math problems step by step", 
    #     language="oromo"
    # )